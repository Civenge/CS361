import math
import socket
import json
import matplotlib.pyplot as plt
from docx import Document

server_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

host = '127.0.0.1'
port = 57318
server_socket.bind((host, port))

server_socket.listen(5)
print(f"Server listening on {host}:{port}")

while True:
    client_socket, addr = server_socket.accept()
    print(f"Connection from {addr}")

    try:
        # receive data in chunks until the entire message is received
        received_data = b""
        while True:
            chunk = client_socket.recv(1024)
            if not chunk:
                print("No chunk received. Breaking out of loop.")
                break
            received_data += chunk
            print(f"Received chunk: {len(chunk)} bytes")

        # decode the received data
        message_in = received_data.decode('utf-8')
        print(f"Message received: {message_in}")

        # parse the JSON data
        data = json.loads(message_in)
        # print(data)

        # remove all NaN "ss" values
        for item in data:
            for key, value in item.items():
                if isinstance(value, float) and math.isnan(value):
                    item[key] = None

        # convert all values to strings
        for item in data:
            for key, value in item.items():
                if value is not None:
                    item[key] = str(value)

        # print(data)
        filtered_data = [entry for entry in data if entry["ss"] is not None]
        # print(filtered_data)

        # extract the relevant information for plotting
        names = [entry["name"] for entry in filtered_data]
        ss_values = [float(entry["ss"]) for entry in filtered_data]

        # create a bar graph
        plt.figure(figsize=(6, 6))
        bars = plt.bar(names, ss_values)

        # rotate x-axis text
        plt.xticks(rotation=45, ha="right")

        plt.xlabel('Names')
        plt.ylabel('SS Values')
        plt.title('SS Values for Different Names')

        # set min value for y-axis
        plt.ylim(70)

        plt.tight_layout()


        # save figure as image
        image_path = 'bar_chart.png'
        plt.savefig(image_path)
        # plt.show()
        # plt.close()

        # create word document
        doc = Document()

        # add heading
        doc.add_heading('Bar Chart of SS Values', level=1)

        # add saved image
        doc.add_picture(image_path)

        # save word document
        doc.save('bar_chart_document.docx')

        # Send the document content back to the client
        with open('bar_chart_document.docx', 'rb') as docx_file:
            while True:
                chunk = docx_file.read(1024)
                if not chunk:
                    break
                client_socket.send(chunk)

        # close connection
        client_socket.close()
        break

    except Exception as e:
        print(f"Error processing message: {e}")
    finally:
        # Close the connection in the finally block to ensure it happens even if an exception occurs
        client_socket.close()
